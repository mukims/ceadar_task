import argparse
import logging
import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Any, List, Optional, Tuple

import warnings

from docx import Document as DocxDocument
from langchain_core.documents import Document
import numpy as np
from openpyxl import load_workbook
from pypdf import PdfReader


logger = logging.getLogger(__name__)

MATHY_LINE = re.compile(
    r"[=+\-*/^_\u2211\u220f\u221a\u2248\u2260\u2264\u2265\u221e\u2208\u2209\u222a\u2229\u2192\u21a6\u2200\u2203]"
    r"|O\([^)]*\)"
)
LATEX_INLINE = re.compile(r"\$[^$]+\$")
DENSE_SYMBOLS = re.compile(
    r"[A-Za-z]\s*\(|\)\s*[A-Za-z]|[A-Za-z]\s*_\s*[A-Za-z0-9]|[A-Za-z]\s*\^\s*[A-Za-z0-9]"
)
FIG_CAPTION_RE = re.compile(r"^(Figure|Fig\.)\s*\d+[\.:]\s+.+", re.IGNORECASE)


def _patch_importlib_metadata() -> None:
    import importlib.metadata as imd

    if getattr(imd.version, "__name__", "") == "_safe_version":
        return

    original_version = imd.version

    def _safe_version(name: str) -> str:
        try:
            return original_version(name)
        except Exception:
            raise imd.PackageNotFoundError(name)

    imd.version = _safe_version


def _import_text_splitter():
    _patch_importlib_metadata()
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        return RecursiveCharacterTextSplitter
    except Exception as exc:
        raise RuntimeError(
            "Failed to import langchain_text_splitters. "
            "This typically means the transformers/huggingface_hub stack is broken.\n"
            "Fix by upgrading the HF stack in your active env:\n"
            "  python -m pip install -U pip setuptools wheel\n"
            "  python -m pip install -U huggingface_hub transformers sentence-transformers\n"
            "If the error persists, run 'python -m pip check' and reinstall the broken package."
        ) from exc


def _import_embeddings():
    _patch_importlib_metadata()
    try:
        from sentence_transformers import SentenceTransformer
        return SentenceTransformer
    except Exception as exc:
        raise RuntimeError(
            "Failed to import SentenceTransformer. "
            "Reinstall or upgrade sentence-transformers in your env."
        ) from exc


def _import_faiss():
    try:
        import faiss  # type: ignore
        return faiss
    except Exception as exc:
        raise RuntimeError(
            "Failed to import faiss. Install faiss-cpu in your env."
        ) from exc


def _import_pipeline():
    _patch_importlib_metadata()
    try:
        from transformers import pipeline
        return pipeline
    except Exception as exc:
        raise RuntimeError(
            "Failed to import transformers.pipeline. "
            "Upgrade/reinstall transformers and huggingface_hub in your active env."
        ) from exc


def _import_pdfplumber(optional: bool = True):
    try:
        import pdfplumber
        return pdfplumber
    except Exception as exc:
        if optional:
            return None
        raise RuntimeError(
            "Failed to import pdfplumber. Install pdfplumber to enable advanced PDF parsing."
        ) from exc


def _import_fitz(optional: bool = True):
    try:
        import fitz  # PyMuPDF
        return fitz
    except Exception as exc:
        if optional:
            return None
        raise RuntimeError(
            "Failed to import PyMuPDF (fitz). Install pymupdf to extract figures from PDFs."
        ) from exc


def _import_pytesseract(optional: bool = True):
    try:
        import pytesseract
        return pytesseract
    except Exception as exc:
        if optional:
            return None
        raise RuntimeError(
            "Failed to import pytesseract. Install pytesseract to OCR figure images."
        ) from exc


def _import_pil_image(optional: bool = True):
    try:
        from PIL import Image
        return Image
    except Exception as exc:
        if optional:
            return None
        raise RuntimeError(
            "Failed to import Pillow. Install pillow to OCR figure images."
        ) from exc


@dataclass(frozen=True)
class RAGConfig:
    files_dir: Path
    persist_dir: Path = Path("faiss_store")
    chunk_size: int = 1000
    chunk_overlap: int = 150
    embedding_model: str = "sentence-transformers/all-MiniLM-L6-v2"
    llm_model: str = "google/flan-t5-base"
    llm_task: Optional[str] = None
    top_k: int = 5
    max_new_tokens: int = 256
    max_context_chars: int = 4000


def _equations(line: str) -> bool:
    line = line.strip()
    if not line or len(line) < 5:
        return False

    score = 0
    if MATHY_LINE.search(line):
        score += 1
    if LATEX_INLINE.search(line):
        score += 1
    if DENSE_SYMBOLS.search(line):
        score += 1

    return score >= 2


def _split_text_in_blocks(text: str) -> List[str]:
    lines = [ln.rstrip() for ln in (text or "").splitlines()]
    blocks: List[str] = []
    buf: List[str] = []

    for ln in lines:
        if ln.strip() == "":
            if buf:
                blocks.append("\n".join(buf).strip())
                buf = []
        else:
            buf.append(ln)

    if buf:
        blocks.append("\n".join(buf).strip())

    return [b for b in blocks if b]


def _multi_line_math(blocks: List[str], max_chars: int = 1500) -> List[Tuple[str, str]]:
    out: List[Tuple[str, str]] = []
    classified: List[Tuple[str, str]] = []

    for block in blocks:
        lines = block.splitlines()
        eq_lines = sum(_equations(ln) for ln in lines)
        if eq_lines >= max(1, len(lines) // 3):
            classified.append(("equation", block))
        else:
            classified.append(("text", block))

    text_buf: List[str] = []
    text_len = 0

    def flush_text() -> None:
        nonlocal text_buf, text_len
        if text_buf:
            merged = "\n\n".join(text_buf).strip()
            start = 0
            while start < len(merged):
                end = min(len(merged), start + max_chars)
                out.append(("text", merged[start:end].strip()))
                start = end
            text_buf = []
            text_len = 0

    for kind, content in classified:
        if kind == "equation":
            flush_text()
            out.append(("equation", content.strip()))
        else:
            if text_len + len(content) + 2 > max_chars:
                flush_text()
            text_buf.append(content)
            text_len += len(content) + 2

    flush_text()
    return out


def _extract_tables_from_page(page) -> List[List[List[str]]]:
    settings = {
        "vertical_strategy": "text",
        "horizontal_strategy": "text",
        "snap_tolerance": 3,
        "intersection_tolerance": 5,
        "join_tolerance": 3,
        "edge_min_length": 3,
        "min_words_vertical": 1,
        "min_words_horizontal": 1,
    }

    tables: List[List[List[str]]] = []
    try:
        table = page.extract_table(table_settings=settings)
        if table and len(table) >= 2:
            tables.append(table)
    except Exception:
        return []
    return tables


def _clean_table(table: List[List[str]]) -> Tuple[List[str], List[List[str]]]:
    def clean_cell(val: Any) -> str:
        return ("" if val is None else str(val)).replace("\n", " ").strip()

    rows = [[clean_cell(c) for c in row] for row in table]
    rows = [row for row in rows if any(c for c in row)]
    if not rows:
        return [], []
    headers = rows[0]
    data = rows[1:]
    headers = [h if h else f"col_{i}" for i, h in enumerate(headers)]
    return headers, data


def _table_to_documents(
    headers: List[str],
    rows: List[List[str]],
    *,
    title: str,
    source: Path,
    page_num: int,
    table_id: str,
) -> List[Document]:
    if not headers:
        return []

    docs: List[Document] = []
    summary = f"Table: {title}\nThis table compares items across columns: {', '.join(headers)}."
    docs.append(
        Document(
            page_content=summary,
            metadata={
                "source": str(source),
                "filename": source.name,
                "filetype": "pdf",
                "page": page_num,
                "table_id": table_id,
                "type": "table_summary",
                "columns": headers,
            },
        )
    )

    for i, row in enumerate(rows):
        row = (row + [""] * len(headers))[: len(headers)]
        claim = (
            f"In the table '{title}', the row item '{row[0]}' has "
            + ", ".join([f"{headers[j]} = {row[j]}" for j in range(1, len(headers))])
            + "."
        )
        docs.append(
            Document(
                page_content=f"Table Row Claim: {claim}",
                metadata={
                    "source": str(source),
                    "filename": source.name,
                    "filetype": "pdf",
                    "page": page_num,
                    "table_id": table_id,
                    "type": "table_row",
                    "row_index": i,
                    "columns": headers,
                    "row_item": row[0],
                },
            )
        )

    return docs


def _figure_captions_from_blocks(blocks: List[str]) -> List[str]:
    captions: List[str] = []
    for block in blocks:
        first = block.splitlines()[0].strip() if block.strip() else ""
        if first and FIG_CAPTION_RE.match(first):
            captions.append(block.strip())
    return captions


def _ocr_image(image_path: Path) -> str:
    pytesseract = _import_pytesseract(optional=True)
    Image = _import_pil_image(optional=True)
    if pytesseract is None or Image is None:
        return ""
    if shutil.which("tesseract") is None:
        return ""

    try:
        img = Image.open(image_path).convert("RGB")
        text = pytesseract.image_to_string(img)
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        return "\n".join(lines)
    except Exception:
        return ""


def _load_pdf_basic(path: Path) -> List[Document]:
    out: List[Document] = []
    reader = PdfReader(str(path))

    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue

        out.append(
            Document(
                page_content=text,
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "filetype": "pdf",
                    "page": i,
                    "type": "text",
                },
            )
        )
    return out


def _load_pdf(path: Path, *, max_text_chars: int = 1200) -> List[Document]:
    pdfplumber = _import_pdfplumber(optional=True)
    if pdfplumber is None:
        logger.warning("pdfplumber not available; falling back to pypdf for %s", path.name)
        return _load_pdf_basic(path)

    fitz = _import_fitz(optional=True)
    fig_doc = None
    if fitz is not None:
        try:
            fig_doc = fitz.open(str(path))
        except Exception:
            fig_doc = None

    out: List[Document] = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            for page_idx, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text() or ""
                blocks = _split_text_in_blocks(page_text)
                typed_blocks = _multi_line_math(blocks, max_chars=max_text_chars)

                for kind, content in typed_blocks:
                    if not content.strip():
                        continue
                    out.append(
                        Document(
                            page_content=content,
                            metadata={
                                "source": str(path),
                                "filename": path.name,
                                "filetype": "pdf",
                                "page": page_idx,
                                "type": kind,
                            },
                        )
                    )

                tables = _extract_tables_from_page(page)
                for t_i, table in enumerate(tables):
                    headers, rows = _clean_table(table)
                    if not headers:
                        continue
                    title = f"Extracted Table (page {page_idx})"
                    table_id = f"p{page_idx}_t{t_i}"
                    out.extend(
                        _table_to_documents(
                            headers,
                            rows,
                            title=title,
                            source=path,
                            page_num=page_idx,
                            table_id=table_id,
                        )
                    )

                for c_i, caption in enumerate(_figure_captions_from_blocks(blocks)):
                    out.append(
                        Document(
                            page_content=f"Figure Caption: {caption}",
                            metadata={
                                "source": str(path),
                                "filename": path.name,
                                "filetype": "pdf",
                                "page": page_idx,
                                "type": "figure_caption",
                                "caption_index": c_i,
                            },
                        )
                    )

                if fig_doc is not None:
                    try:
                        fig_page = fig_doc[page_idx - 1]
                    except Exception:
                        fig_page = None

                    if fig_page is not None:
                        images = fig_page.get_images(full=True)
                        for img_i, img in enumerate(images):
                            xref = img[0]
                            try:
                                base = fig_doc.extract_image(xref)
                            except Exception:
                                continue
                            img_bytes = base.get("image")
                            if not img_bytes:
                                continue
                            ext = base.get("ext", "png")

                            out_dir = Path("out/figures")
                            out_dir.mkdir(parents=True, exist_ok=True)
                            img_path = out_dir / f"page{page_idx:03d}_fig{img_i:02d}.{ext}"
                            try:
                                img_path.write_bytes(img_bytes)
                            except Exception:
                                continue

                            ocr_text = _ocr_image(img_path)
                            if ocr_text.strip():
                                out.append(
                                    Document(
                                        page_content=f"Figure OCR Text:\n{ocr_text}",
                                        metadata={
                                            "source": str(path),
                                            "filename": path.name,
                                            "filetype": "pdf",
                                            "page": page_idx,
                                            "type": "figure_ocr",
                                            "figure_index": img_i,
                                            "image_path": str(img_path),
                                        },
                                    )
                                )
    except Exception as exc:
        logger.warning("pdfplumber failed for %s (%s); using pypdf fallback.", path.name, exc)
        return _load_pdf_basic(path)
    finally:
        if fig_doc is not None:
            try:
                fig_doc.close()
            except Exception:
                pass

    return out


def _load_docx(path: Path) -> List[Document]:
    out: List[Document] = []
    doc = DocxDocument(str(path))

    for pi, para in enumerate(doc.paragraphs, start=1):
        text = (para.text or "").strip()
        if not text:
            continue
        out.append(
            Document(
                page_content=text,
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "filetype": "docx",
                    "block": f"paragraph:{pi}",
                },
            )
        )

    for ti, table in enumerate(doc.tables, start=1):
        rows_text: List[str] = []
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            if any(cells):
                rows_text.append("\t".join(cells))

        table_text = "\n".join(rows_text).strip()
        if table_text:
            out.append(
                Document(
                    page_content=table_text,
                    metadata={
                        "source": str(path),
                        "filename": path.name,
                        "filetype": "docx",
                        "block": f"table:{ti}",
                    },
                )
            )

    return out


def _load_xlsx(path: Path) -> List[Document]:
    out: List[Document] = []
    wb = load_workbook(filename=str(path), read_only=True, data_only=True)

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]

        lines: List[str] = []
        for row in ws.iter_rows(values_only=True):
            row_vals = [("" if c is None else str(c)).strip() for c in row]
            if not any(row_vals):
                continue
            lines.append("\t".join(row_vals))

        sheet_text = "\n".join(lines).strip()
        if sheet_text:
            out.append(
                Document(
                    page_content=sheet_text,
                    metadata={
                        "source": str(path),
                        "filename": path.name,
                        "filetype": "xlsx",
                        "sheet": sheet_name,
                    },
                )
            )

    return out


def load_documents(files_dir: Path, *, pdf_max_text_chars: int = 1200) -> List[Document]:
    files_dir = Path(files_dir)
    if not files_dir.exists():
        raise FileNotFoundError(f"Folder not found: {files_dir}")

    docs: List[Document] = []
    files_required = sorted(
        p for p in files_dir.rglob("*") if p.is_file() and p.suffix.lower() in {".pdf", ".docx", ".xlsx"}
    )

    if not files_required:
        raise FileNotFoundError("No supported files found (.pdf, .docx, .xlsx).")

    for path in files_required:
        ext = path.suffix.lower()
        try:
            if ext == ".pdf":
                docs.extend(_load_pdf(path, max_text_chars=pdf_max_text_chars))
            elif ext == ".docx":
                docs.extend(_load_docx(path))
            elif ext == ".xlsx":
                docs.extend(_load_xlsx(path))
        except Exception as exc:
            logger.exception("Skipping %s: %s", path, exc)

    return docs


def chunk_documents(cfg: RAGConfig, docs: List[Document]) -> List[Document]:
    RecursiveCharacterTextSplitter = _import_text_splitter()
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=cfg.chunk_size,
        chunk_overlap=cfg.chunk_overlap,
    )
    return splitter.split_documents(docs)


def _faiss_index_path(persist_dir: Path) -> Path:
    persist_dir.mkdir(parents=True, exist_ok=True)
    return persist_dir / "faiss.index"


def _faiss_docs_path(persist_dir: Path) -> Path:
    persist_dir.mkdir(parents=True, exist_ok=True)
    return persist_dir / "docs.jsonl"


def _l2_normalize(vectors: np.ndarray) -> np.ndarray:
    norms = np.linalg.norm(vectors, axis=1, keepdims=True)
    norms[norms == 0] = 1.0
    return vectors / norms


def _encode_chunks(cfg: RAGConfig, chunks: List[Document]) -> np.ndarray:
    SentenceTransformer = _import_embeddings()
    model = SentenceTransformer(cfg.embedding_model)
    texts = [doc.page_content for doc in chunks]
    embeddings = model.encode(texts, convert_to_numpy=True, show_progress_bar=True)
    return _l2_normalize(embeddings.astype(np.float32))


def _save_docs(path: Path, chunks: List[Document]) -> None:
    import json

    with path.open("w", encoding="utf-8") as f:
        for doc in chunks:
            f.write(
                json.dumps(
                    {"page_content": doc.page_content, "metadata": doc.metadata},
                    ensure_ascii=True,
                )
                + "\n"
            )


def _load_docs(path: Path) -> List[Document]:
    import json

    docs: List[Document] = []
    with path.open("r", encoding="utf-8") as f:
        for line in f:
            data = json.loads(line)
            docs.append(Document(page_content=data["page_content"], metadata=data["metadata"]))
    return docs


def build_faiss_index(cfg: RAGConfig, chunks: List[Document]):
    faiss = _import_faiss()
    persist_dir = Path(cfg.persist_dir)
    index_path = _faiss_index_path(persist_dir)
    docs_path = _faiss_docs_path(persist_dir)

    if index_path.exists() and docs_path.exists():
        index = faiss.read_index(str(index_path))
        docs = _load_docs(docs_path)
        return index, docs

    embeddings = _encode_chunks(cfg, chunks)
    index = faiss.IndexFlatIP(embeddings.shape[1])
    index.add(embeddings)
    faiss.write_index(index, str(index_path))
    _save_docs(docs_path, chunks)
    return index, chunks


def retrieve_documents(index, docs: List[Document], cfg: RAGConfig, query: str, top_k: int) -> List[Document]:
    SentenceTransformer = _import_embeddings()
    model = SentenceTransformer(cfg.embedding_model)
    query_emb = model.encode([query], convert_to_numpy=True)
    query_emb = _l2_normalize(query_emb.astype(np.float32))

    distances, indices = index.search(query_emb, top_k)
    results: List[Document] = []
    for idx in indices[0]:
        if idx < 0 or idx >= len(docs):
            continue
        results.append(docs[idx])
    return results


def _build_prompt(
    query: str,
    docs: List[Document],
    max_context_chars: int,
    tokenizer=None,
    max_input_tokens: Optional[int] = None,
) -> str:
    prefix = (
        "Answer the question using only the context below. "
        "If the answer is not in the context, say you don't know.\n\n"
        "Context:\n"
    )
    suffix = f"\n\nQuestion: {query}\nAnswer:"

    parts: List[str] = []
    total_chars = 0

    if tokenizer and max_input_tokens:
        base_tokens = len(tokenizer.encode(prefix + suffix, add_special_tokens=False))
        budget = max_input_tokens - base_tokens
        if budget < 16:
            budget = 16
        used = 0
        for i, doc in enumerate(docs, start=1):
            snippet = doc.page_content.strip()
            if not snippet:
                continue
            entry = f"[{i}] {snippet}"
            entry_tokens = len(tokenizer.encode(entry, add_special_tokens=False))
            if used + entry_tokens > budget:
                break
            parts.append(entry)
            used += entry_tokens
    else:
        for i, doc in enumerate(docs, start=1):
            snippet = doc.page_content.strip()
            if not snippet:
                continue
            entry = f"[{i}] {snippet}"
            if total_chars + len(entry) > max_context_chars:
                break
            parts.append(entry)
            total_chars += len(entry)

    context = "\n\n".join(parts)
    return f"{prefix}{context}{suffix}"


def _infer_llm_task(model_name: str) -> str:
    model_name = model_name.lower()
    if "t5" in model_name or "flan" in model_name:
        return "text2text-generation"
    return "text-generation"


def generate_answer(cfg: RAGConfig, query: str, docs: List[Document]) -> str:
    task = cfg.llm_task or _infer_llm_task(cfg.llm_model)
    pipeline = _import_pipeline()
    generator = pipeline(task, model=cfg.llm_model)
    max_len = getattr(generator.tokenizer, "model_max_length", 512)
    if not isinstance(max_len, int) or max_len > 100000:
        max_len = 512
    input_budget = max_len - cfg.max_new_tokens
    if input_budget < 16:
        input_budget = max_len

    prompt = _build_prompt(
        query,
        docs,
        cfg.max_context_chars,
        tokenizer=generator.tokenizer,
        max_input_tokens=input_budget,
    )

    outputs = generator(
        prompt,
        max_new_tokens=cfg.max_new_tokens,
        do_sample=False,
        num_return_sequences=1,
        eos_token_id=getattr(generator.tokenizer, "eos_token_id", None),
    )
    generated = outputs[0].get("generated_text", "").strip()
    if task == "text-generation" and generated.startswith(prompt):
        return generated[len(prompt):].strip()
    return generated


def run_pipeline(cfg: RAGConfig, query: str) -> None:
    logger.info("Loading documents from %s", cfg.files_dir)
    raw_docs = load_documents(cfg.files_dir)

    logger.info("Chunking %d documents", len(raw_docs))
    chunks = chunk_documents(cfg, raw_docs)

    logger.info("Building/Loading FAISS index at %s", cfg.persist_dir)
    index, docs = build_faiss_index(cfg, chunks)

    top_k = min(cfg.top_k, 5)
    logger.info("Retrieving top %d chunks for query", top_k)
    results = retrieve_documents(index, docs, cfg, query, top_k)

    seen = set()
    unique_results: List[Document] = []
    for doc in results:
        key = doc.page_content.strip().lower()
        if not key or key in seen:
            continue
        seen.add(key)
        unique_results.append(doc)
        if len(unique_results) >= top_k:
            break

    print("\nTop retrieved chunks (max 5, de-duplicated):\n")
    for i, doc in enumerate(unique_results, start=1):
        meta = doc.metadata
        preview = doc.page_content[:200].replace("\n", " ")
        print(f"[{i}] {meta.get('filename', meta.get('source', 'unknown'))} | {preview}...")

    print("\nGenerated answer:\n")
    answer = generate_answer(cfg, query, unique_results)
    print(answer)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Single-file RAG pipeline")
    parser.add_argument("--files-dir", type=Path, default=Path("."), help="Folder to scan for documents")
    parser.add_argument("--persist-dir", type=Path, default=Path("faiss_store"), help="FAISS persistence directory")
    parser.add_argument("--query", type=str, required=True, help="User query")
    parser.add_argument("--chunk-size", type=int, default=1000)
    parser.add_argument("--chunk-overlap", type=int, default=150)
    parser.add_argument("--embedding-model", type=str, default="sentence-transformers/all-MiniLM-L6-v2")
    parser.add_argument("--llm-model", type=str, default="google/flan-t5-base")
    parser.add_argument("--llm-task", type=str, default=None, help="Override pipeline task (e.g., text-generation)")
    parser.add_argument("--top-k", type=int, default=5)
    parser.add_argument("--max-new-tokens", type=int, default=256)
    parser.add_argument("--max-context-chars", type=int, default=4000)
    return parser.parse_args()


def main() -> None:
    warnings.filterwarnings("ignore")
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
    logging.getLogger("pypdf").setLevel(logging.ERROR)
    try:
        from transformers import logging as hf_logging

        hf_logging.set_verbosity_error()
    except Exception:
        pass
    args = parse_args()

    cfg = RAGConfig(
        files_dir=args.files_dir,
        persist_dir=args.persist_dir,
        chunk_size=args.chunk_size,
        chunk_overlap=args.chunk_overlap,
        embedding_model=args.embedding_model,
        llm_model=args.llm_model,
        llm_task=args.llm_task,
        top_k=args.top_k,
        max_new_tokens=args.max_new_tokens,
        max_context_chars=args.max_context_chars,
    )

    run_pipeline(cfg, args.query)


if __name__ == "__main__":
    main()
