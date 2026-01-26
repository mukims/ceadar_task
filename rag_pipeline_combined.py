"""
Extended RAG Pipeline
=====================

This module extends the original ``rag_pipeline`` by incorporating
additional PDF parsing capabilities inspired by the provided Jupyter
notebook.  In addition to simple text extraction, PDFs are parsed for
mathematical expressions, tables, figure captions and figure images
with optional OCR.  The resulting chunks are stored as
``langchain_core.documents.Document`` instances with rich metadata
describing their origin and type (text, equation, table summary,
table row, figure caption or figure OCR).

The extended extraction functions are purely additive – if the
optional dependencies used for table extraction (``pdfplumber``) or
image OCR (``pytesseract``) are missing, the pipeline still works,
falling back gracefully to simple text and math classification.

To use this module, run it similarly to the original ``rag_pipeline``:

.. code-block:: console

   python rag_pipeline_combined.py --files-dir ./documents --query "What is the energy equation?"

This will load all PDFs, Word and Excel files from ``files-dir``,
build a FAISS index from the resulting chunks and answer the provided
query using a language model.
"""

import argparse
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple

import warnings

# Optional import for Word documents.  If python-docx is unavailable
# (common in minimal environments), DocxDocument will remain ``None`` and
# DOCX files will be skipped gracefully.
try:
    from docx import Document as DocxDocument  # type: ignore
except Exception:
    DocxDocument = None  # type: ignore
# Attempt to import the Document class from langchain_core.  If the
# package is unavailable (common in minimal environments), define a
# simple substitute with ``page_content`` and ``metadata`` attributes.
try:
    from langchain_core.documents import Document  # type: ignore
except Exception:
    from dataclasses import dataclass as _dc
    from typing import Any as _Any, Dict as _Dict
    @_dc
    class Document:
        """Fallback Document type used when langchain_core is not installed.

        This simple dataclass mimics the interface of ``langchain_core.documents.Document``
        by storing a text payload in ``page_content`` and arbitrary metadata in
        ``metadata``.  Any code expecting the real Document class should
        continue to work with this substitute.
        """
        page_content: str
        metadata: _Dict[str, _Any]
import numpy as np
from openpyxl import load_workbook
# Attempt to import a PDF reader.  ``pypdf`` is preferred but may not be
# installed in the execution environment.  As a fallback we try
# ``PyPDF2``.  If neither is available, ``PdfReader`` will remain ``None``
# and the pipeline will attempt to extract text using PyMuPDF instead.
try:
    from pypdf import PdfReader  # type: ignore
except Exception:
    try:
        from PyPDF2 import PdfReader  # type: ignore
    except Exception:
        PdfReader = None  # type: ignore


logger = logging.getLogger(__name__)


def _patch_importlib_metadata() -> None:
    import importlib.metadata as imd

    if getattr(imd.version, "__name__", "") == "_safe_version":
        return

    original_version = imd.version

    def _safe_version(name: str) -> str:
        try:
            return original_version(name)
        except Exception:
            raise imd.PackageNotFoundError(name)

    imd.version = _safe_version


def _import_text_splitter():
    _patch_importlib_metadata()
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        return RecursiveCharacterTextSplitter
    except Exception as exc:
        raise RuntimeError(
            "Failed to import langchain_text_splitters. "
            "This typically means the transformers/huggingface_hub stack is broken.\n"
            "Fix by upgrading the HF stack in your active env:\n"
            "  python -m pip install -U pip setuptools wheel\n"
            "  python -m pip install -U huggingface_hub transformers sentence-transformers\n"
            "If the error persists, run 'python -m pip check' and reinstall the broken package."
        ) from exc


def _import_embeddings():
    _patch_importlib_metadata()
    try:
        from sentence_transformers import SentenceTransformer
        return SentenceTransformer
    except Exception as exc:
        raise RuntimeError(
            "Failed to import SentenceTransformer. "
            "Reinstall or upgrade sentence-transformers in your env."
        ) from exc


def _import_faiss():
    try:
        import faiss  # type: ignore
        return faiss
    except Exception as exc:
        raise RuntimeError(
            "Failed to import faiss. Install faiss-cpu in your env."
        ) from exc


def _import_pipeline():
    _patch_importlib_metadata()
    try:
        from transformers import pipeline
        return pipeline
    except Exception as exc:
        raise RuntimeError(
            "Failed to import transformers.pipeline. "
            "Upgrade/reinstall transformers and huggingface_hub in your active env."
        ) from exc


@dataclass(frozen=True)
class RAGConfig:
    files_dir: Path
    persist_dir: Path = Path("faiss_store")
    chunk_size: int = 1000
    chunk_overlap: int = 150
    embedding_model: str = "sentence-transformers/all-MiniLM-L6-v2"
    llm_model: str = "google/flan-t5-base"
    llm_task: Optional[str] = None
    top_k: int = 5
    max_new_tokens: int = 256
    max_context_chars: int = 4000


# ---------------------------------------------------------------------------
# Utility functions for PDF parsing
#
# The following helpers are adapted from the provided Jupyter notebook. They
# classify mathematical lines, split text into blocks, detect multi‑line
# equations, extract tables and figure captions and perform simple OCR on
# extracted images. Where external dependencies such as pdfplumber or
# pytesseract are unavailable, the functions degrade gracefully (returning
# empty results rather than raising exceptions). These helpers are only used
# when loading PDFs; other file types are handled as in the original
# pipeline.

import re


# Regular expressions for identifying mathematical content
MATHY_LINE = re.compile(r"[=+\-*/^_∑∏√≈≠≤≥∞∈∉∪∩→↦∀∃]|O\([^)]*\)")
LATEX_INLINE = re.compile(r"\$[^$]+\$")  # inline LaTeX delimited by $...$
DENSE_SYMBOLS = re.compile(r"[A-Za-z]\s*\(|\)\s*[A-Za-z]|[A-Za-z]\s*_\s*[A-Za-z0-9]|[A-Za-z]\s*\^\s*[A-Za-z0-9]")


def equations(line: str) -> bool:
    """Heuristically determine if a line contains mathematical content.

    A line is considered mathematical if it is non‑empty, reasonably long
    and contains at least two indicators such as operators, LaTeX inline
    fragments or dense symbol patterns.
    """
    line = line.strip()
    if not line:
        return False
    if len(line) < 5:
        return False
    score = 0
    if MATHY_LINE.search(line):
        score += 1
    if LATEX_INLINE.search(line):
        score += 1
    if DENSE_SYMBOLS.search(line):
        score += 1
    return score >= 2


def split_text_in_blocks(text: str) -> List[str]:
    """Split a string into paragraphs separated by blank lines."""
    lines = [ln.rstrip() for ln in (text or "").splitlines()]
    blocks: List[str] = []
    buf: List[str] = []
    for ln in lines:
        if ln.strip() == "":
            if buf:
                blocks.append("\n".join(buf).strip())
                buf = []
        else:
            buf.append(ln)
    if buf:
        blocks.append("\n".join(buf).strip())
    return [b for b in blocks if b]


def multi_line_math(blocks: List[str], max_chars: int = 1500) -> List[Tuple[str, str]]:
    """Classify blocks as text or equations and chunk long text.

    Blocks classified as equations are emitted as‑is.  Text blocks are
    concatenated until ``max_chars`` is reached to form chunks suitable
    for embedding.
    """
    out: List[Tuple[str, str]] = []
    classified: List[Tuple[str, str]] = []
    # classify each block
    for b in blocks:
        lines = b.splitlines()
        eq_lines = sum(equations(ln) for ln in lines)
        if eq_lines >= max(1, len(lines) // 3):
            classified.append(("equation", b))
        else:
            classified.append(("text", b))
    # accumulate text blocks up to max_chars
    text_buf: List[str] = []
    text_len = 0
    def flush_text() -> None:
        nonlocal text_buf, text_len
        if text_buf:
            merged = "\n\n".join(text_buf).strip()
            start = 0
            while start < len(merged):
                end = min(len(merged), start + max_chars)
                out.append(("text", merged[start:end].strip()))
                start = end
            text_buf = []
            text_len = 0
    for kind, content in classified:
        if kind == "equation":
            flush_text()
            out.append(("equation", content.strip()))
        else:
            if text_len + len(content) + 2 > max_chars:
                flush_text()
            text_buf.append(content)
            text_len += len(content) + 2
    flush_text()
    return out


def extract_tables_from_page(page) -> List[List[List[str]]]:
    """Attempt to extract tables from a pdfplumber page.

    Returns a list of tables, each represented as list of rows, each row
    being a list of cell strings.  If pdfplumber is not available or
    extraction fails, an empty list is returned.
    """
    try:
        # Avoid importing pdfplumber at module load time since it may not
        # be installed.  Import locally to catch ImportError gracefully.
        import pdfplumber  # type: ignore
    except Exception:
        return []
    # table extraction settings tuned for borderless tables
    settings = {
        "vertical_strategy": "text",
        "horizontal_strategy": "text",
        "snap_tolerance": 3,
        "intersection_tolerance": 5,
        "join_tolerance": 3,
        "edge_min_length": 3,
        "min_words_vertical": 1,
        "min_words_horizontal": 1,
    }
    tables: List[List[List[str]]] = []
    try:
        table = page.extract_table(table_settings=settings)
        if table and len(table) >= 2:
            tables.append(table)
    except Exception:
        pass
    return tables


def clean_table(table: List[List[str]]) -> Tuple[List[str], List[List[str]]]:
    """Basic cleanup: strip cells and drop empty rows.

    Assumes the first row contains headers.  Empty header cells are
    replaced with generic column names.  Returns a tuple of headers and
    data rows.
    """
    def clean_cell(x: Any) -> str:
        return (x or "").replace("\n", " ").strip()
    rows = [[clean_cell(c) for c in r] for r in table]
    rows = [r for r in rows if any(c for c in r)]
    if not rows:
        return [], []
    headers = rows[0]
    data = rows[1:]
    headers = [h if h else f"col_{i}" for i, h in enumerate(headers)]
    return headers, data


def table_to_docs(headers: List[str], rows: List[List[str]], *, title: str,
                  source: str, page_num: int, table_id: str) -> List[Document]:
    """Create summary and row documents from a table.

    Produces a summary document describing the table and a document for
    each row that reads like a natural language claim.  Metadata
    includes the original source filename, page number and table ID.
    """
    docs: List[Document] = []
    summary_text = (
        f"Table: {title}\n"
        f"This table compares items across columns: {', '.join(headers)}."
    )
    docs.append(
        Document(
            page_content=summary_text,
            metadata={
                "source": source,
                "page": page_num,
                "table_id": table_id,
                "type": "table_summary",
                "columns": headers,
            },
        )
    )
    for i, row in enumerate(rows):
        # pad short rows
        row = (row + [""] * len(headers))[: len(headers)]
        claim = (
            f"In the table '{title}', the row item '{row[0]}' has "
            + ", ".join([f"{headers[j]} = {row[j]}" for j in range(1, len(headers))])
            + "."
        )
        docs.append(
            Document(
                page_content=f"Table Row Claim: {claim}",
                metadata={
                    "source": source,
                    "page": page_num,
                    "table_id": table_id,
                    "type": "table_row",
                    "row_index": i,
                    "columns": headers,
                    "row_item": row[0],
                },
            )
        )
    return docs


FIG_CAPTION_RE = re.compile(r"^(Figure|Fig\.)\s*\d+[\.:]\s+.+", re.IGNORECASE)


def figure_captions_from_blocks(blocks: List[str]) -> List[str]:
    """Return lines that look like figure captions from a list of blocks."""
    caps: List[str] = []
    for b in blocks:
        first = b.splitlines()[0].strip()
        if FIG_CAPTION_RE.match(first):
            caps.append(b.strip())
    return caps


def ocr_image(image_path: Path) -> str:
    """Perform OCR on an image if ``pytesseract`` is available.

    Returns extracted text as a single string with lines joined by newlines.
    If the Tesseract binary or the ``pytesseract`` module is unavailable
    the function returns an empty string.
    """
    try:
        import shutil
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except Exception:
        return ""
    # ensure tesseract binary exists in PATH
    if shutil.which("tesseract") is None:
        return ""
    try:
        img = Image.open(image_path).convert("RGB")
        text = pytesseract.image_to_string(img)
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        return "\n".join(lines)
    except Exception:
        return ""


def ingest_pdf_to_documents(pdf_path: Path, max_text_chars: int = 1500) -> List[Document]:
    """Parse a PDF into a list of Documents with rich metadata.

    The PDF is processed page by page.  For each page, the following
    information is extracted:

    * Text and equations: the page text is split into blocks and each
      block is classified as either ordinary text or an equation.  Text
      blocks are concatenated up to ``max_text_chars`` characters to
      produce manageable chunks.  Equations are kept separate.  Each
      chunk yields a Document with type ``text`` or ``equation``.

    * Tables: if pdfplumber is available, tables are extracted using
      borderless detection.  A summary document and one document per row
      are produced for each table.

    * Figure captions: simple regex matching is used to extract figure
      captions from the page text.  Each caption yields a Document.

    * Figure images and OCR: using PyMuPDF (``fitz``) any images on the
      page are extracted to an ``out/figures`` directory relative to
      the working directory.  If Tesseract OCR is available, the
      extracted image is OCRed and the resulting text yields a
      ``figure_ocr`` Document.  The image file is saved regardless of
      OCR availability.

    If pdfplumber or PyMuPDF are unavailable, the respective parts of
    extraction are skipped.  The function returns a list of Documents
    representing all extracted content.
    """
    all_docs: List[Document] = []
    source = pdf_path.name
    # Use PyMuPDF to access images; if unavailable skip figure extraction
    try:
        import fitz  # type: ignore
        fitz_doc = fitz.open(str(pdf_path))
    except Exception:
        fitz_doc = None
    # Use pdfplumber for text and tables if available
    try:
        import pdfplumber  # type: ignore
        plumber_doc = pdfplumber.open(str(pdf_path))
    except Exception:
        plumber_doc = None
    # Determine number of pages and text extraction method.  Prefer
    # PdfReader if available, otherwise fall back to PyMuPDF (fitz) for
    # text extraction.  If both are unavailable, no text will be
    # extracted.
    num_pages = 0
    reader = None
    if PdfReader is not None:
        try:
            reader = PdfReader(str(pdf_path))
            num_pages = len(reader.pages)
        except Exception:
            reader = None
    if num_pages == 0 and fitz_doc is not None:
        try:
            num_pages = fitz_doc.page_count  # type: ignore[attr-defined]
        except Exception:
            num_pages = 0
    # iterate through pages (1-indexed for metadata)
    for page_idx in range(1, (num_pages or 0) + 1):
        # -----------------------------------------------------------------
        # 1) TEXT AND EQUATIONS
        # -----------------------------------------------------------------
        page_text = ""
        # try to use pypdf/PyPDF2 reader
        if reader is not None and page_idx - 1 < len(getattr(reader, "pages", [])):
            try:
                pg = reader.pages[page_idx - 1]
                # some PdfReader implementations expose extract_text differently
                page_text = getattr(pg, "extract_text", lambda: "")() or ""
            except Exception:
                page_text = ""
        # fallback to fitz for text
        if not page_text and fitz_doc is not None:
            try:
                pg = fitz_doc[page_idx - 1]
                page_text = pg.get_text("text") or ""
            except Exception:
                page_text = ""
        # split and classify
        blocks = split_text_in_blocks(page_text)
        typed_blocks = multi_line_math(blocks, max_chars=max_text_chars)
        for kind, content in typed_blocks:
            all_docs.append(
                Document(
                    page_content=content,
                    metadata={
                        "source": source,
                        "page": page_idx,
                        "type": kind,
                    },
                )
            )
        # -----------------------------------------------------------------
        # 2) TABLES (pdfplumber only)
        # -----------------------------------------------------------------
        if plumber_doc is not None:
            try:
                page_plumber = plumber_doc.pages[page_idx - 1]
                tables = extract_tables_from_page(page_plumber)
            except Exception:
                tables = []
            for t_i, tbl in enumerate(tables):
                headers, rows = clean_table(tbl)
                if not headers:
                    continue
                title = f"Extracted Table (page {page_idx})"
                table_id = f"p{page_idx}_t{t_i}"
                all_docs.extend(
                    table_to_docs(
                        headers,
                        rows,
                        title=title,
                        source=source,
                        page_num=page_idx,
                        table_id=table_id,
                    )
                )
        # -----------------------------------------------------------------
        # 3) FIGURE CAPTIONS
        # -----------------------------------------------------------------
        caps = figure_captions_from_blocks(blocks)
        for c_i, caption in enumerate(caps):
            all_docs.append(
                Document(
                    page_content=f"Figure Caption: {caption}",
                    metadata={
                        "source": source,
                        "page": page_idx,
                        "type": "figure_caption",
                        "caption_index": c_i,
                    },
                )
            )
        # -----------------------------------------------------------------
        # 4) FIGURE IMAGES AND OCR (PyMuPDF only)
        # -----------------------------------------------------------------
        if fitz_doc is not None:
            try:
                fig_page = fitz_doc[page_idx - 1]
                images = fig_page.get_images(full=True)
            except Exception:
                images = []
            for img_i, img in enumerate(images):
                try:
                    xref = img[0]
                    base = fitz_doc.extract_image(xref)
                    img_bytes = base["image"]
                    ext = base.get("ext", "png")
                except Exception:
                    continue
                out_dir = Path("out/figures")
                out_dir.mkdir(parents=True, exist_ok=True)
                img_path = out_dir / f"{source}_page{page_idx:03d}_fig{img_i:02d}.{ext}"
                try:
                    img_path.write_bytes(img_bytes)
                except Exception:
                    pass
                # attempt OCR on the saved image
                ocr_text = ocr_image(img_path)
                if ocr_text.strip():
                    all_docs.append(
                        Document(
                            page_content=f"Figure OCR Text:\n{ocr_text}",
                            metadata={
                                "source": source,
                                "page": page_idx,
                                "type": "figure_ocr",
                                "figure_index": img_i,
                                "image_path": str(img_path),
                            },
                        )
                    )
    if plumber_doc is not None:
        try:
            plumber_doc.close()
        except Exception:
            pass
    if fitz_doc is not None:
        try:
            fitz_doc.close()
        except Exception:
            pass
    return all_docs


# ---------------------------------------------------------------------------
# Loading functions for different file types
#
# These mirror the original helpers but with PDF ingestion extended to use
# the rich parsing implemented above.

def _load_pdf(path: Path) -> List[Document]:
    """Load a PDF into Documents using the extended ingestion routine."""
    try:
        docs = ingest_pdf_to_documents(path)
        return docs
    except Exception as exc:
        logger.exception("Failed to ingest PDF with extended parser: %s", exc)
        # Fallback: attempt minimal extraction using whatever PDF reader is available.
        out: List[Document] = []
        # Try PdfReader first
        if PdfReader is not None:
            try:
                reader = PdfReader(str(path))
                for i, page in enumerate(getattr(reader, "pages", []), start=1):
                    try:
                        text = getattr(page, "extract_text", lambda: "")() or ""
                    except Exception:
                        text = ""
                    if not text:
                        continue
                    out.append(
                        Document(
                            page_content=text,
                            metadata={
                                "source": str(path),
                                "filename": path.name,
                                "filetype": "pdf",
                                "page": i,
                            },
                        )
                    )
                return out
            except Exception:
                pass
        # Fall back to PyMuPDF if available
        try:
            import fitz  # type: ignore
            doc = fitz.open(str(path))
            for i in range(doc.page_count):
                try:
                    text = doc[i].get_text("text") or ""
                except Exception:
                    text = ""
                if not text.strip():
                    continue
                out.append(
                    Document(
                        page_content=text.strip(),
                        metadata={
                            "source": str(path),
                            "filename": path.name,
                            "filetype": "pdf",
                            "page": i + 1,
                        },
                    )
                )
            doc.close()
        except Exception:
            pass
        return out


def _load_docx(path: Path) -> List[Document]:
    """Load paragraphs and tables from a Word document.

    If python-docx is unavailable, this function returns an empty list
    and logs a message.  Each paragraph and each table becomes a
    separate Document with associated metadata.
    """
    if DocxDocument is None:
        logger.warning(
            "python-docx is not installed; skipping DOCX file: %s", path
        )
        return []
    out: List[Document] = []
    try:
        doc = DocxDocument(str(path))
    except Exception as exc:
        logger.exception("Failed to read DOCX %s: %s", path, exc)
        return []
    # paragraphs
    for pi, para in enumerate(doc.paragraphs, start=1):
        text = (para.text or "").strip()
        if not text:
            continue
        out.append(
            Document(
                page_content=text,
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "filetype": "docx",
                    "block": f"paragraph:{pi}",
                },
            )
        )
    # tables
    for ti, table in enumerate(doc.tables, start=1):
        rows_text: List[str] = []
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            if any(cells):
                rows_text.append("\t".join(cells))
        table_text = "\n".join(rows_text).strip()
        if table_text:
            out.append(
                Document(
                    page_content=table_text,
                    metadata={
                        "source": str(path),
                        "filename": path.name,
                        "filetype": "docx",
                        "block": f"table:{ti}",
                    },
                )
            )
    return out


def _load_xlsx(path: Path) -> List[Document]:
    out: List[Document] = []
    wb = load_workbook(filename=str(path), read_only=True, data_only=True)
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines: List[str] = []
        for row in ws.iter_rows(values_only=True):
            row_vals = [(("" if c is None else str(c)).strip()) for c in row]
            if not any(row_vals):
                continue
            lines.append("\t".join(row_vals))
        sheet_text = "\n".join(lines).strip()
        if sheet_text:
            out.append(
                Document(
                    page_content=sheet_text,
                    metadata={
                        "source": str(path),
                        "filename": path.name,
                        "filetype": "xlsx",
                        "sheet": sheet_name,
                    },
                )
            )
    return out


def load_documents(files_dir: Path) -> List[Document]:
    files_dir = Path(files_dir)
    if not files_dir.exists():
        raise FileNotFoundError(f"Folder not found: {files_dir}")
    docs: List[Document] = []
    files_required = sorted(
        p for p in files_dir.rglob("*") if p.is_file() and p.suffix.lower() in {".pdf", ".docx", ".xlsx"}
    )
    if not files_required:
        raise FileNotFoundError("No supported files found (.pdf, .docx, .xlsx).")
    for path in files_required:
        ext = path.suffix.lower()
        try:
            if ext == ".pdf":
                docs.extend(_load_pdf(path))
            elif ext == ".docx":
                docs.extend(_load_docx(path))
            elif ext == ".xlsx":
                docs.extend(_load_xlsx(path))
        except Exception as exc:
            logger.exception("Skipping %s: %s", path, exc)
    return docs


def chunk_documents(cfg: RAGConfig, docs: List[Document]) -> List[Document]:
    RecursiveCharacterTextSplitter = _import_text_splitter()
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=cfg.chunk_size,
        chunk_overlap=cfg.chunk_overlap,
    )
    return splitter.split_documents(docs)


def _faiss_index_path(persist_dir: Path) -> Path:
    persist_dir.mkdir(parents=True, exist_ok=True)
    return persist_dir / "faiss.index"


def _faiss_docs_path(persist_dir: Path) -> Path:
    persist_dir.mkdir(parents=True, exist_ok=True)
    return persist_dir / "docs.jsonl"


def _l2_normalize(vectors: np.ndarray) -> np.ndarray:
    norms = np.linalg.norm(vectors, axis=1, keepdims=True)
    norms[norms == 0] = 1.0
    return vectors / norms


def _encode_chunks(cfg: RAGConfig, chunks: List[Document]) -> np.ndarray:
    SentenceTransformer = _import_embeddings()
    model = SentenceTransformer(cfg.embedding_model)
    texts = [doc.page_content for doc in chunks]
    embeddings = model.encode(texts, convert_to_numpy=True, show_progress_bar=True)
    return _l2_normalize(embeddings.astype(np.float32))


def _save_docs(path: Path, chunks: List[Document]) -> None:
    import json
    with path.open("w", encoding="utf-8") as f:
        for doc in chunks:
            f.write(
                json.dumps(
                    {"page_content": doc.page_content, "metadata": doc.metadata},
                    ensure_ascii=True,
                )
                + "\n"
            )


def _load_docs(path: Path) -> List[Document]:
    import json
    docs: List[Document] = []
    with path.open("r", encoding="utf-8") as f:
        for line in f:
            data = json.loads(line)
            docs.append(Document(page_content=data["page_content"], metadata=data["metadata"]))
    return docs


def build_faiss_index(cfg: RAGConfig, chunks: List[Document]):
    faiss = _import_faiss()
    persist_dir = Path(cfg.persist_dir)
    index_path = _faiss_index_path(persist_dir)
    docs_path = _faiss_docs_path(persist_dir)
    if index_path.exists() and docs_path.exists():
        index = faiss.read_index(str(index_path))
        docs = _load_docs(docs_path)
        return index, docs
    embeddings = _encode_chunks(cfg, chunks)
    index = faiss.IndexFlatIP(embeddings.shape[1])
    index.add(embeddings)
    faiss.write_index(index, str(index_path))
    _save_docs(docs_path, chunks)
    return index, chunks


def retrieve_documents(index, docs: List[Document], cfg: RAGConfig, query: str, top_k: int) -> List[Document]:
    SentenceTransformer = _import_embeddings()
    model = SentenceTransformer(cfg.embedding_model)
    query_emb = model.encode([query], convert_to_numpy=True)
    query_emb = _l2_normalize(query_emb.astype(np.float32))
    distances, indices = index.search(query_emb, top_k)
    results: List[Document] = []
    for idx in indices[0]:
        if idx < 0 or idx >= len(docs):
            continue
        results.append(docs[idx])
    return results


def _build_prompt(
    query: str,
    docs: List[Document],
    max_context_chars: int,
    tokenizer=None,
    max_input_tokens: Optional[int] = None,
) -> str:
    prefix = (
        "Answer the question using only the context below. "
        "If the answer is not in the context, say you don't know.\n\n"
        "Context:\n"
    )
    suffix = f"\n\nQuestion: {query}\nAnswer:"
    parts: List[str] = []
    total_chars = 0
    if tokenizer and max_input_tokens:
        base_tokens = len(tokenizer.encode(prefix + suffix, add_special_tokens=False))
        budget = max_input_tokens - base_tokens
        if budget < 16:
            budget = 16
        used = 0
        for i, doc in enumerate(docs, start=1):
            snippet = doc.page_content.strip()
            if not snippet:
                continue
            entry = f"[{i}] {snippet}"
            entry_tokens = len(tokenizer.encode(entry, add_special_tokens=False))
            if used + entry_tokens > budget:
                break
            parts.append(entry)
            used += entry_tokens
    else:
        for i, doc in enumerate(docs, start=1):
            snippet = doc.page_content.strip()
            if not snippet:
                continue
            entry = f"[{i}] {snippet}"
            if total_chars + len(entry) > max_context_chars:
                break
            parts.append(entry)
            total_chars += len(entry)
    context = "\n\n".join(parts)
    return f"{prefix}{context}{suffix}"


def _infer_llm_task(model_name: str) -> str:
    model_name = model_name.lower()
    if "t5" in model_name or "flan" in model_name:
        return "text2text-generation"
    return "text-generation"


def generate_answer(cfg: RAGConfig, query: str, docs: List[Document]) -> str:
    task = cfg.llm_task or _infer_llm_task(cfg.llm_model)
    pipeline = _import_pipeline()
    try:
        generator = pipeline(task, model=cfg.llm_model)
    except KeyError:
        task = "text-generation"
        generator = pipeline(task, model=cfg.llm_model)
    max_len = getattr(generator.tokenizer, "model_max_length", 512)
    if not isinstance(max_len, int) or max_len > 100000:
        max_len = 512
    input_budget = max_len - cfg.max_new_tokens
    if input_budget < 16:
        input_budget = max_len
    prompt = _build_prompt(
        query,
        docs,
        cfg.max_context_chars,
        tokenizer=generator.tokenizer,
        max_input_tokens=input_budget,
    )
    outputs = generator(
        prompt,
        max_new_tokens=cfg.max_new_tokens,
        do_sample=False,
        num_return_sequences=1,
        eos_token_id=getattr(generator.tokenizer, "eos_token_id", None),
    )
    generated = outputs[0].get("generated_text", "").strip()
    if task == "text-generation" and generated.startswith(prompt):
        return generated[len(prompt):].strip()
    return generated


def run_pipeline(cfg: RAGConfig, query: str) -> None:
    logger.info("Loading documents from %s", cfg.files_dir)
    raw_docs = load_documents(cfg.files_dir)
    logger.info("Chunking %d documents", len(raw_docs))
    chunks = chunk_documents(cfg, raw_docs)
    logger.info("Building/Loading FAISS index at %s", cfg.persist_dir)
    index, docs = build_faiss_index(cfg, chunks)
    top_k = min(cfg.top_k, 5)
    logger.info("Retrieving top %d chunks for query", top_k)
    results = retrieve_documents(index, docs, cfg, query, top_k)
    seen = set()
    unique_results: List[Document] = []
    for doc in results:
        key = doc.page_content.strip().lower()
        if not key or key in seen:
            continue
        seen.add(key)
        unique_results.append(doc)
        if len(unique_results) >= top_k:
            break
    print("\nTop retrieved chunks (max 5, de‑duplicated):\n")
    for i, doc in enumerate(unique_results, start=1):
        meta = doc.metadata
        preview = doc.page_content[:200].replace("\n", " ")
        print(f"[{i}] {meta.get('filename', meta.get('source', 'unknown'))} | {preview}...")
    print("\nGenerated answer:\n")
    answer = generate_answer(cfg, query, unique_results)
    print(answer)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Single‑file RAG pipeline with extended PDF parsing")
    parser.add_argument("--files-dir", type=Path, default=Path("."), help="Folder to scan for documents")
    parser.add_argument("--persist-dir", type=Path, default=Path("faiss_store"), help="FAISS persistence directory")
    parser.add_argument("--query", type=str, required=True, help="User query")
    parser.add_argument("--chunk-size", type=int, default=1000)
    parser.add_argument("--chunk-overlap", type=int, default=150)
    parser.add_argument(
        "--embedding-model",
        type=str,
        default="sentence-transformers/all-MiniLM-L6-v2",
        help="Name of the embedding model to use (HuggingFace repo ID)",
    )
    parser.add_argument("--llm-model", type=str, default="google/flan-t5-base")
    parser.add_argument("--llm-task", type=str, default=None, help="Override pipeline task (e.g., text-generation)")
    parser.add_argument("--top-k", type=int, default=5)
    parser.add_argument("--max-new-tokens", type=int, default=256)
    parser.add_argument("--max-context-chars", type=int, default=4000)
    return parser.parse_args()


def main() -> None:
    warnings.filterwarnings("ignore")
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
    logging.getLogger("pypdf").setLevel(logging.ERROR)
    try:
        from transformers import logging as hf_logging  # type: ignore
        hf_logging.set_verbosity_error()
    except Exception:
        pass
    args = parse_args()
    cfg = RAGConfig(
        files_dir=args.files_dir,
        persist_dir=args.persist_dir,
        chunk_size=args.chunk_size,
        chunk_overlap=args.chunk_overlap,
        embedding_model=args.embedding_model,
        llm_model=args.llm_model,
        llm_task=args.llm_task,
        top_k=args.top_k,
        max_new_tokens=args.max_new_tokens,
        max_context_chars=args.max_context_chars,
    )
    run_pipeline(cfg, args.query)


if __name__ == "__main__":
    main()
